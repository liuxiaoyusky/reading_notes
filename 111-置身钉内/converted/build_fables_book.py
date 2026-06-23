"""
Build the combined fables + source book for `111-置身钉内`.

Layout per chapter: 寓言 (chapter title + fable) → 原文 (all sections in that chapter).
This pairs the distilled fable (the mirror image) with the original long-form
reflection, so a reader can hold both side by side.

Inputs:
    - converted/fables/<chapter>/<NN>-<chapter-title>.md   (9 fables)
    - converted/sections/<chapter>/*.md                     (9 chapters × many sections)
Outputs (written to converted/):
    - fables-book.md    merged single-file Markdown
    - fables-book.docx  Word document with proper Chinese rendering
    - fables-book.pdf   PDF with proper Chinese rendering
"""

import os
import re
import datetime
from pathlib import Path

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
HERE = Path(__file__).resolve().parent
FABLES = HERE / "fables"
SECTIONS = HERE / "sections"
OUT_MD = HERE / "fables-book.md"
OUT_DOCX = HERE / "fables-book.docx"
OUT_PDF = HERE / "fables-book.pdf"

# ---------------------------------------------------------------------------
# Chapter metadata — must match fables/ and sections/ subfolder names exactly.
# Fable file stem is "01-<chapter-title>.md" (Chinese), so numbering is positional
# (one fable per chapter, index = chapter order).
# ---------------------------------------------------------------------------
CHAPTERS = [
    ("00-序章", "序章"),
    ("01-发心第一", "发心第一"),
    ("02-定位第二", "定位第二"),
    ("03-设计第三", "设计第三"),
    ("04-用户第四", "用户第四"),
    ("05-敏捷第五", "敏捷第五"),
    ("06-秩序第六", "秩序第六"),
    ("07-军争第七", "军争第七"),
    ("08-长期第八", "长期第八"),
]

BOOK_TITLE = "置身钉内"
BOOK_SUBTITLE = "寓言 + 原文 · 双面合订"
AUTHOR_ORIGINAL = "原作：个九（钉钉 ONE 产品经理）"
ATTRIBUTION = "寓言整理：Claude · 9 章 · 9 篇寓言"
TODAY = datetime.date.today().isoformat()

# CJK font candidates (macOS) — first one that registers wins.
CJK_FONT_CANDIDATES = [
    ("STHeiti", "/System/Library/Fonts/STHeiti Medium.ttc", 0),
    ("STHeitiLight", "/System/Library/Fonts/STHeiti Light.ttc", 0),
    ("PingFang", "/Library/Fonts/Arial Unicode.ttf", 0),  # not CJK but last-ditch fallback
]

# ---------------------------------------------------------------------------
# Walking & parsing
# ---------------------------------------------------------------------------

def chapter_dir(slug: str) -> Path:
    return FABLES / slug


def fable_files_in_chapter(slug: str):
    """Glob fable files in a chapter. Each chapter has exactly 1 fable file
    named '01-<chapter-title>.md'. We sort by the leading NN digits."""
    d = chapter_dir(slug)
    files = []
    for p in d.glob("*.md"):
        # Extract leading NN from stem (e.g. "01-发心第一" -> 1).
        m = re.match(r'^(\d+)-', p.stem)
        if m:
            files.append((int(m.group(1)), p))
    files.sort(key=lambda x: x[0])
    return [p for _, p in files]


def section_files_in_chapter(slug: str):
    """Glob section files in converted/sections/<slug>/, sorted by NN prefix.
    Skip 00_index.md (just navigation)."""
    d = SECTIONS / slug
    files = []
    for p in d.glob("*.md"):
        if p.stem == "00_index":
            continue
        m = re.match(r'^(\d+)-', p.stem)
        if m:
            files.append((int(m.group(1)), p))
    files.sort(key=lambda x: x[0])
    return [p for _, p in files]


def parse_section(path: Path) -> dict:
    """Parse a section .md (原始原文). Strips YAML frontmatter and returns
    the body as plain markdown text. The '## 节名' heading stays as H2;
    the renderer will downgrade to H3 when emitting."""
    text = path.read_text(encoding="utf-8")
    lines = text.splitlines()
    # Strip YAML frontmatter
    if lines and lines[0].strip() == "---":
        for i in range(1, len(lines)):
            if lines[i].strip() == "---":
                lines = lines[i + 1:]
                break
    body = "\n".join(lines).strip()
    return {"title": path.stem, "body": body}


def parse_fable(path: Path) -> dict:
    """
    Parse one fable .md file. Structure (verified across all 101 files):

        # 标题
        [body — may contain `---` scene breaks, and one 📝 来源 line]
        ---
        ## 原文定义
        > quote lines...
        ## 对应点
        | 故事元素 | 概念对应 |
        |---|---|
        | row 1   | row 1   |
        ...

    Returns dict with keys: title, body, quote_lines, table_rows, source_line.
    """
    text = path.read_text(encoding="utf-8")
    lines = text.splitlines()

    # Strip YAML frontmatter (--- ... ---) at the very top if present.
    if lines and lines[0].strip() == "---":
        for i in range(1, len(lines)):
            if lines[i].strip() == "---":
                lines = lines[i + 1:]
                break

    # Title = first H1.
    title = ""
    for ln in lines:
        m = re.match(r"^#\s+(.+?)\s*$", ln)
        if m:
            title = m.group(1).strip()
            break
    # Fallback: derive a placeholder title from the file stem (e.g. "04-04" -> "寓言 4-4").
    if not title:
        title = f"寓言 {path.stem}"

    # Locate the `## 原文定义` and `## 对应点` H2 line numbers.
    # 111- fables use bold text `**📖 原文定义**` instead of H2 — accept both.
    i_quote = i_table = None
    for i, ln in enumerate(lines):
        s = ln.strip()
        if i_quote is None and (s == "## 原文定义" or s == "**📖 原文定义**" or "原文定义" in s and s.startswith("**")):
            i_quote = i
        elif i_table is None and (s == "## 对应点" or s == "**💡 对应点**" or "对应点" in s and s.startswith("**")):
            i_table = i
    if i_quote is None or i_table is None:
        raise ValueError(f"{path}: missing 原文定义 or 对应点 section")

    # Body region: from after the first H1 to before `## 原文定义`.
    # If there is no H1, start from line 0.
    body_start = 0
    for i, ln in enumerate(lines):
        if re.match(r"^#\s+", ln):
            body_start = i + 1
            break
    body_lines = lines[body_start:i_quote]
    # Strip leading blank lines, trailing blank lines.
    while body_lines and not body_lines[0].strip():
        body_lines.pop(0)
    while body_lines and not body_lines[-1].strip():
        body_lines.pop()

    # Extract source line — the (only) line containing 📝 来源 in the whole file.
    # The line may or may not be prefixed with `> ` (blockquote). Store WITHOUT the 📝 prefix.
    source_line = ""
    for ln in lines:
        if "📝 来源" in ln:
            cleaned = re.sub(r"^>\s*", "", ln).strip()
            cleaned = re.sub(r"^📝\s*", "", cleaned).strip()
            source_line = cleaned
            break

    # Remove the source line from the body region (it lives inside body, just before ---).
    body_lines = [ln for ln in body_lines if "📝 来源" not in ln]
    body = "\n".join(body_lines).rstrip()

    # Quote region: between `## 原文定义` and `## 对应点`.
    quote_region = lines[i_quote + 1:i_table]
    quote_lines = []
    for ln in quote_region:
        stripped = ln.strip()
        if not stripped:
            continue
        # Strip leading `> ` and `>` markers; keep inner text only.
        cleaned = re.sub(r"^>\s?", "", ln).rstrip()
        quote_lines.append(cleaned)

    # Table region: after `## 对应点`.
    table_region = lines[i_table + 1:]
    table_rows = []
    for ln in table_region:
        s = ln.strip()
        if not s.startswith("|"):
            continue
        if s.count("|") < 3:
            continue
        cells = [c.strip() for c in s.strip("|").split("|")]
        # Skip the separator row `|---|---|`
        if all(re.fullmatch(r"[-:\s]+", c) for c in cells):
            continue
        table_rows.append(cells)
    if table_rows and len(table_rows[0]) == 2 and table_rows[0][0] == "故事元素":
        # The renderer adds its own header; drop the parsed one to avoid duplication.
        table_rows = table_rows[1:]

    return {
        "title": title,
        "body": body,
        "quote_lines": quote_lines,
        "table_rows": table_rows,
        "source_line": source_line,
    }


# ---------------------------------------------------------------------------
# 1. Merged Markdown
# ---------------------------------------------------------------------------

def write_merged_markdown(fables_by_chapter: dict) -> int:
    """
    Write the merged fables-book.md. Returns the line count.
    fables_by_chapter: {chapter_slug: [(section_no, parsed_fable, source_path), ...]}
    """
    total_fables = sum(len(v) for v in fables_by_chapter.values())

    out = []
    out.append(f"# {BOOK_TITLE}\n")
    out.append(f"## {BOOK_SUBTITLE}\n")
    out.append(f"{AUTHOR_ORIGINAL}\n")
    out.append(f"{ATTRIBUTION}\n")
    out.append(f"生成日期：{TODAY}\n")
    out.append(f"寓言数量：{total_fables} 篇 · 共 {len(CHAPTERS)} 章\n")
    out.append("")

    out.append("## 目录\n")
    for idx, (slug, title) in enumerate(CHAPTERS, 1):
        n = len(fables_by_chapter[slug])
        ns = len(fables_by_chapter[slug + "_sections"])
        out.append(f"- 第 {idx} 章 {title} （{n} 篇寓言 · {ns} 节原文）")
    out.append("")

    out.append("---\n")

    for idx, (slug, title) in enumerate(CHAPTERS, 1):
        out.append(f"# 第 {idx} 章 {title}\n")
        n = len(fables_by_chapter[slug])
        ns = len(fables_by_chapter[slug + "_sections"])
        out.append(f"> 本章包含 {n} 篇寓言 · {ns} 节原文\n")
        out.append("")

        # --- 寓言 part ---
        out.append(f"## 寓言\n")
        for sec_no, fable, _src in fables_by_chapter[slug]:
            out.append(f"### {fable['title']}\n")
            out.append(fable["body"])
            out.append("")
            out.append("#### 原文定义\n")
            for q in fable["quote_lines"]:
                out.append(f"> {q}")
            out.append("")
            out.append("#### 对应点\n")
            table_rows = fable["table_rows"]
            if table_rows and len(table_rows[0]) == 2 and table_rows[0][0] == "故事元素":
                table_rows = table_rows[1:]
            if table_rows:
                out.append("| 故事元素 | 概念对应 |")
                out.append("|---|---|")
                for row in table_rows:
                    cells = (row + ["", ""])[:2]
                    out.append(f"| {cells[0]} | {cells[1]} |")
            else:
                out.append("_(无)_")
            out.append("")
            if fable["source_line"]:
                out.append(f"📝 {fable['source_line']}")
                out.append("")
            out.append("---\n")

        # --- 原文 part ---
        out.append(f"## 原文\n")
        for sec in fables_by_chapter[slug + "_sections"]:
            # Each section body is plain markdown with an H2 '## 节名' header.
            # Demote that H2 to H3 so it sits under '## 原文'.
            body = sec["body"]
            body_demoted = re.sub(r"^##\s+", "### ", body, count=1, flags=re.MULTILINE)
            out.append(body_demoted)
            out.append("")
            out.append("---\n")

    md_text = "\n".join(out)
    OUT_MD.write_text(md_text, encoding="utf-8")
    return md_text.count("\n") + 1


# ---------------------------------------------------------------------------
# 2. DOCX
# ---------------------------------------------------------------------------

def _set_cjk(run, name: str, size: float, bold: bool = False, italic: bool = False, color=None):
    """Set both Latin and East-Asian font on a run so CJK glyphs render correctly."""
    from docx.shared import Pt, RGBColor
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)


def _add_page_number_footer(section, cjk_name: str):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Clear any existing content
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    run = p.add_run()
    fldChar1 = OxmlElement("w:fldChar"); fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText"); instrText.text = "PAGE   \\* MERGEFORMAT"
    fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2)
    _set_cjk(run, cjk_name, size=9, color=None)


def build_docx(fables_by_chapter: dict, sections_by_chapter: dict, cjk_name: str):
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    total_fables = sum(len(v) for v in fables_by_chapter.values())

    doc = Document()

    # Default body style: CJK-aware.
    style = doc.styles["Normal"]
    style.font.name = cjk_name
    style.font.size = Pt(10.5)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), cjk_name)
    rFonts.set(qn("w:ascii"), cjk_name)
    rFonts.set(qn("w:hAnsi"), cjk_name)

    # Page setup: A4 with 2.5cm margins.
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ----- Title page -----
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    r = p.add_run(BOOK_TITLE)
    _set_cjk(r, cjk_name, size=28, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(BOOK_SUBTITLE)
    _set_cjk(r, cjk_name, size=14, italic=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    r = p.add_run(AUTHOR_ORIGINAL)
    _set_cjk(r, cjk_name, size=12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(ATTRIBUTION)
    _set_cjk(r, cjk_name, size=12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(80)
    r = p.add_run(f"生成日期：{TODAY}")
    _set_cjk(r, cjk_name, size=10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"寓言数量：{total_fables} 篇 · 共 {len(CHAPTERS)} 章")
    _set_cjk(r, cjk_name, size=10)

    # Page break.
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)

    # ----- Table of contents -----
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("目录")
    _set_cjk(r, cjk_name, size=22, bold=True)
    p.paragraph_format.space_after = Pt(18)

    for idx, (slug, en_title) in enumerate(CHAPTERS, 1):
        n = len(fables_by_chapter[slug])
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(f"第 {idx} 章　{en_title}　（{n} 篇）")
        _set_cjk(r, cjk_name, size=12)

    # Page break before chapters.
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)

    # ----- Chapters & fables -----
    for idx, (slug, en_title) in enumerate(CHAPTERS, 1):
        n = len(fables_by_chapter[slug])

        # Chapter cover page.
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(180)
        r = p.add_run(f"第 {idx} 章")
        _set_cjk(r, cjk_name, size=18, bold=False)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(18)
        r = p.add_run(en_title)
        _set_cjk(r, cjk_name, size=24, bold=True)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(40)
        r = p.add_run(f"本章包含 {n} 篇寓言")
        _set_cjk(r, cjk_name, size=12, italic=True)

        # Page break to start fables.
        p = doc.add_paragraph()
        p.add_run().add_break(WD_BREAK.PAGE)

        for sec_no, fable, _src in fables_by_chapter[slug]:
            # Fable H2 title.
            h = doc.add_heading(level=2)
            h.paragraph_format.space_before = Pt(6)
            h.paragraph_format.space_after = Pt(12)
            h.paragraph_format.keep_with_next = True
            r = h.add_run(fable["title"])
            _set_cjk(r, cjk_name, size=18, bold=True)
            # Make sure the heading style itself uses CJK font for the anchor numbering.
            hPr = h._element.get_or_add_pPr()
            h_rPr = OxmlElement("w:rPr")
            h_rFonts = OxmlElement("w:rFonts")
            h_rFonts.set(qn("w:eastAsia"), cjk_name)
            h_rFonts.set(qn("w:ascii"), cjk_name)
            h_rFonts.set(qn("w:hAnsi"), cjk_name)
            h_rPr.append(h_rFonts)
            hPr.append(h_rPr)

            # Body — split on blank lines (paragraphs).
            for para in re.split(r"\n\s*\n", fable["body"]):
                para = para.strip()
                if not para:
                    continue
                # Handle in-body `---` as a centered separator.
                if re.fullmatch(r"-{3,}", para):
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r = p.add_run("· · ·")
                    _set_cjk(r, cjk_name, size=10, color=RGBColor(0x99, 0x99, 0x99))
                    continue
                # Handle stray H2 markers like `## 故事` — render as italic subheading.
                if para.startswith("## "):
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(8)
                    p.paragraph_format.space_after = Pt(4)
                    r = p.add_run(para[3:].strip())
                    _set_cjk(r, cjk_name, size=12, bold=True, italic=True)
                    continue
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Cm(0.74)  # 2 chars
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.line_spacing = 1.5
                # Convert intra-paragraph newlines to soft line breaks.
                lines = para.split("\n")
                for j, ln in enumerate(lines):
                    if j > 0:
                        p.add_run().add_break()
                    r = p.add_run(ln)
                    _set_cjk(r, cjk_name, size=10.5)

            # Spacer before next section.
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)

            # `## 原文定义` heading.
            h = doc.add_heading(level=3)
            r = h.add_run("原文定义")
            _set_cjk(r, cjk_name, size=13, bold=True)
            # Force CJK on the heading style.
            hPr = h._element.get_or_add_pPr()
            h_rPr = OxmlElement("w:rPr")
            h_rFonts = OxmlElement("w:rFonts")
            h_rFonts.set(qn("w:eastAsia"), cjk_name)
            h_rFonts.set(qn("w:ascii"), cjk_name)
            h_rFonts.set(qn("w:hAnsi"), cjk_name)
            h_rPr.append(h_rFonts)
            hPr.append(h_rPr)

            for q in fable["quote_lines"]:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.8)
                p.paragraph_format.right_indent = Cm(0.8)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.4
                r = p.add_run(q)
                _set_cjk(r, cjk_name, size=10.5, italic=True,
                         color=RGBColor(0x40, 0x40, 0x40))

            # `## 对应点` heading.
            h = doc.add_heading(level=3)
            r = h.add_run("对应点")
            _set_cjk(r, cjk_name, size=13, bold=True)
            hPr = h._element.get_or_add_pPr()
            h_rPr = OxmlElement("w:rPr")
            h_rFonts = OxmlElement("w:rFonts")
            h_rFonts.set(qn("w:eastAsia"), cjk_name)
            h_rFonts.set(qn("w:ascii"), cjk_name)
            h_rFonts.set(qn("w:hAnsi"), cjk_name)
            h_rPr.append(h_rFonts)
            hPr.append(h_rPr)

            if fable["table_rows"]:
                # 2-column table.
                tbl = doc.add_table(rows=1 + len(fable["table_rows"]), cols=2)
                tbl.style = "Table Grid"
                # Column widths.
                widths = [Cm(7.5), Cm(8.0)]
                # Header.
                hdr = tbl.rows[0].cells
                hdr[0].text = ""
                hdr[1].text = ""
                for cell, text in zip(hdr, ["故事元素", "概念对应"]):
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r = p.add_run(text)
                    _set_cjk(r, cjk_name, size=10.5, bold=True)
                    # Shade header.
                    tcPr = cell._tc.get_or_add_tcPr()
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:val"), "clear")
                    shd.set(qn("w:color"), "auto")
                    shd.set(qn("w:fill"), "D9D9D9")
                    tcPr.append(shd)
                # Body rows.
                for ri, row in enumerate(fable["table_rows"], start=1):
                    cells = (row + ["", ""])[:2]
                    for ci, val in enumerate(cells):
                        cell = tbl.rows[ri].cells[ci]
                        p = cell.paragraphs[0]
                        p.paragraph_format.line_spacing = 1.3
                        r = p.add_run(val)
                        _set_cjk(r, cjk_name, size=10)
                    tbl.rows[ri].height = Cm(0.6)
                for row in tbl.rows:
                    for ci, w in enumerate(widths):
                        row.cells[ci].width = w
            else:
                p = doc.add_paragraph()
                r = p.add_run("（无）")
                _set_cjk(r, cjk_name, size=10, italic=True,
                         color=RGBColor(0x80, 0x80, 0x80))

            # Source line.
            if fable["source_line"]:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(8)
                r = p.add_run(f"📝 {fable['source_line']}")
                _set_cjk(r, cjk_name, size=9, color=RGBColor(0x80, 0x80, 0x80))

            # Bottom separator.
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(18)
            r = p.add_run("· · ·")
            _set_cjk(r, cjk_name, size=10, color=RGBColor(0x99, 0x99, 0x99))

        # ----- 原文 part: render all sections of this chapter -----
        chapter_sections = sections_by_chapter.get(slug, [])
        if chapter_sections:
            # Section divider page.
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(80)
            r = p.add_run("—— 原文 ——")
            _set_cjk(r, cjk_name, size=14, italic=True,
                     color=RGBColor(0x66, 0x66, 0x66))
            p.add_run().add_break(WD_BREAK.PAGE)

            for sec in chapter_sections:
                # Section body has '## 节名' as its first H2; render as H3 in docx.
                body_lines = sec["body"].split("\n")
                for ln_idx, para in enumerate(re.split(r"\n\s*\n", "\n".join(body_lines))):
                    para = para.strip()
                    if not para:
                        continue
                    if re.fullmatch(r"-{3,}", para):
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        r = p.add_run("· · ·")
                        _set_cjk(r, cjk_name, size=9, color=RGBColor(0x99, 0x99, 0x99))
                        continue
                    if para.startswith("## "):
                        # Demote H2 to H3 for source sections.
                        h = doc.add_heading(level=3)
                        h.paragraph_format.space_before = Pt(14)
                        h.paragraph_format.space_after = Pt(6)
                        h.paragraph_format.keep_with_next = True
                        r = h.add_run(para[3:].strip())
                        _set_cjk(r, cjk_name, size=13, bold=True)
                        hPr = h._element.get_or_add_pPr()
                        h_rPr = OxmlElement("w:rPr")
                        h_rFonts = OxmlElement("w:rFonts")
                        h_rFonts.set(qn("w:eastAsia"), cjk_name)
                        h_rFonts.set(qn("w:ascii"), cjk_name)
                        h_rFonts.set(qn("w:hAnsi"), cjk_name)
                        h_rPr.append(h_rFonts)
                        hPr.append(h_rPr)
                        continue
                    if para.startswith("!["):
                        # Image — skip in print book (we don't have docx image relinking).
                        continue
                    p = doc.add_paragraph()
                    p.paragraph_format.first_line_indent = Cm(0.74)  # 2 chars
                    p.paragraph_format.space_after = Pt(3)
                    p.paragraph_format.line_spacing = 1.5
                    for j, ln in enumerate(para.split("\n")):
                        if j > 0:
                            p.add_run().add_break()
                        r = p.add_run(ln)
                        _set_cjk(r, cjk_name, size=10)

    # Page-number footer.
    for sec in doc.sections:
        _add_page_number_footer(sec, cjk_name)

    doc.save(str(OUT_DOCX))


# ---------------------------------------------------------------------------
# 3. PDF
# ---------------------------------------------------------------------------

def _register_cjk_font() -> str:
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    for name, path, sub_idx in CJK_FONT_CANDIDATES:
        if not os.path.exists(path):
            continue
        try:
            pdfmetrics.registerFont(TTFont(name, path, subfontIndex=sub_idx))
            print(f"[pdf] using CJK font: {name} from {path}")
            return name
        except Exception as e:
            print(f"[pdf] failed to register {name} from {path}: {e}")
    raise RuntimeError("No usable CJK font found on this system.")


def _escape_xml(s: str) -> str:
    return (s.replace("&", "&amp;")
             .replace("<", "&lt;")
             .replace(">", "&gt;"))


def build_pdf(fables_by_chapter: dict, sections_by_chapter: dict, cjk_name: str):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.lib import colors
    from reportlab.platypus import (BaseDocTemplate, PageTemplate, Frame,
                                    Paragraph, Spacer, PageBreak, Table, TableStyle,
                                    NextPageTemplate)

    total_fables = sum(len(v) for v in fables_by_chapter.values())
    page_w, page_h = A4

    # Two templates: cover/separator pages (no page number) and content pages (with page number).
    def cover_page(canv, doc):
        canv.saveState()
        canv.setFont(cjk_name, 9)
        canv.setFillColor(colors.grey)
        # No page number on cover.
        canv.restoreState()

    def content_page(canv, doc):
        canv.saveState()
        canv.setFont(cjk_name, 9)
        canv.setFillColor(colors.grey)
        canv.drawCentredString(page_w / 2, 1.2 * cm, str(canv.getPageNumber()))
        canv.restoreState()

    doc = BaseDocTemplate(
        str(OUT_PDF),
        pagesize=A4,
        leftMargin=2.5 * cm, rightMargin=2.5 * cm,
        topMargin=2.5 * cm, bottomMargin=2.5 * cm,
        title=BOOK_TITLE,
        author=ATTRIBUTION,
    )

    frame_cover = Frame(0, 0, page_w, page_h, id="cover",
                        leftPadding=2.5 * cm, rightPadding=2.5 * cm,
                        topPadding=2.5 * cm, bottomPadding=2.5 * cm)
    frame_content = Frame(2.5 * cm, 2.5 * cm, page_w - 5 * cm, page_h - 5 * cm,
                          id="content", leftPadding=0, rightPadding=0,
                          topPadding=0, bottomPadding=0)

    doc.addPageTemplates([
        PageTemplate(id="cover", frames=[frame_cover], onPage=cover_page),
        PageTemplate(id="content", frames=[frame_content], onPage=content_page),
    ])

    # ----- Styles -----
    s_title = ParagraphStyle("title", fontName=cjk_name, fontSize=32, leading=44,
                             alignment=TA_CENTER, spaceBefore=120, spaceAfter=12)
    s_subtitle = ParagraphStyle("subtitle", fontName=cjk_name, fontSize=15, leading=22,
                                alignment=TA_CENTER, spaceAfter=40,
                                textColor=colors.HexColor("#444444"))
    s_meta = ParagraphStyle("meta", fontName=cjk_name, fontSize=11, leading=18,
                            alignment=TA_CENTER, spaceAfter=4)
    s_meta_lg = ParagraphStyle("metalarge", fontName=cjk_name, fontSize=12, leading=20,
                               alignment=TA_CENTER, spaceAfter=6)
    s_toc_h = ParagraphStyle("toch", fontName=cjk_name, fontSize=22, leading=30,
                             spaceBefore=0, spaceAfter=20)
    s_toc_item = ParagraphStyle("tocitem", fontName=cjk_name, fontSize=12, leading=22,
                                spaceAfter=6)
    s_chap_num = ParagraphStyle("chapnum", fontName=cjk_name, fontSize=18, leading=28,
                                alignment=TA_CENTER, spaceBefore=180, spaceAfter=12)
    s_chap_title = ParagraphStyle("chaptitle", fontName=cjk_name, fontSize=26, leading=40,
                                  alignment=TA_CENTER, spaceAfter=30, bold=True)
    s_chap_meta = ParagraphStyle("chapmeta", fontName=cjk_name, fontSize=12, leading=20,
                                 alignment=TA_CENTER, spaceAfter=0,
                                 textColor=colors.HexColor("#666666"))
    s_fable_title = ParagraphStyle("fabletitle", fontName=cjk_name, fontSize=18, leading=28,
                                   spaceBefore=8, spaceAfter=12, bold=True)
    s_body = ParagraphStyle("body", fontName=cjk_name, fontSize=10.5, leading=18,
                            firstLineIndent=21, spaceAfter=4, alignment=TA_LEFT)
    s_body_noindent = ParagraphStyle("bodynoindent", fontName=cjk_name, fontSize=10.5,
                                     leading=18, spaceAfter=4, alignment=TA_LEFT)
    s_h2_inline = ParagraphStyle("h2inline", fontName=cjk_name, fontSize=12, leading=20,
                                 spaceBefore=10, spaceAfter=6, bold=True, italic=True)
    s_sep = ParagraphStyle("sep", fontName=cjk_name, fontSize=10, leading=14,
                           alignment=TA_CENTER, spaceBefore=4, spaceAfter=4,
                           textColor=colors.HexColor("#999999"))
    s_h3 = ParagraphStyle("h3", fontName=cjk_name, fontSize=13, leading=22,
                          spaceBefore=12, spaceAfter=6, bold=True)
    s_quote = ParagraphStyle("quote", fontName=cjk_name, fontSize=10.5, leading=18,
                             leftIndent=20, rightIndent=20, spaceAfter=2,
                             textColor=colors.HexColor("#404040"), italic=True)
    s_source = ParagraphStyle("source", fontName=cjk_name, fontSize=9, leading=14,
                              alignment=TA_CENTER, spaceBefore=8, spaceAfter=4,
                              textColor=colors.HexColor("#808080"))

    story = []

    # ----- Title page (cover) -----
    story.append(Paragraph(_escape_xml(BOOK_TITLE), s_title))
    story.append(Paragraph(_escape_xml(BOOK_SUBTITLE), s_subtitle))
    story.append(Spacer(1, 40))
    story.append(Paragraph(_escape_xml(AUTHOR_ORIGINAL), s_meta_lg))
    story.append(Paragraph(_escape_xml(ATTRIBUTION), s_meta_lg))
    story.append(Spacer(1, 60))
    story.append(Paragraph(f"生成日期：{_escape_xml(TODAY)}", s_meta))
    story.append(Paragraph(f"寓言数量：{total_fables} 篇 · 共 {len(CHAPTERS)} 章", s_meta))

    # Switch to content template (with page numbers) starting next page.
    story.append(NextPageTemplate("content"))
    story.append(PageBreak())

    # ----- Table of contents -----
    story.append(Paragraph("目录", s_toc_h))
    for idx, (slug, en_title) in enumerate(CHAPTERS, 1):
        n = len(fables_by_chapter[slug])
        story.append(Paragraph(
            f"第 {idx} 章　{_escape_xml(en_title)}　（{n} 篇）", s_toc_item))
    story.append(PageBreak())

    # ----- Chapters -----
    for idx, (slug, en_title) in enumerate(CHAPTERS, 1):
        n = len(fables_by_chapter[slug])

        # Chapter cover (back to "cover" template? No — keep numbers. Just a big page.)
        story.append(Spacer(1, 60))
        story.append(Paragraph(f"第 {idx} 章", s_chap_num))
        story.append(Paragraph(_escape_xml(en_title), s_chap_title))
        story.append(Spacer(1, 30))
        story.append(Paragraph(f"本章包含 {n} 篇寓言", s_chap_meta))
        story.append(PageBreak())

        for sec_no, fable, _src in fables_by_chapter[slug]:
            story.append(Paragraph(_escape_xml(fable["title"]), s_fable_title))
            for para in re.split(r"\n\s*\n", fable["body"]):
                para = para.strip()
                if not para:
                    continue
                if re.fullmatch(r"-{3,}", para):
                    story.append(Paragraph("· · ·", s_sep))
                    continue
                if para.startswith("## "):
                    story.append(Paragraph(_escape_xml(para[3:].strip()), s_h2_inline))
                    continue
                # Convert intra-paragraph newlines to <br/>.
                story.append(Paragraph(_escape_xml(para).replace("\n", "<br/>"), s_body))

            story.append(Spacer(1, 4))
            story.append(Paragraph("原文定义", s_h3))
            quote_html = "<br/>".join(_escape_xml(q) for q in fable["quote_lines"])
            if quote_html:
                story.append(Paragraph(quote_html, s_quote))

            story.append(Paragraph("对应点", s_h3))
            if fable["table_rows"]:
                # Wrap each cell in a Paragraph so long Chinese text wraps instead of overflowing.
                cell_style = ParagraphStyle(
                    "tablecell", fontName=cjk_name, fontSize=9.5, leading=14,
                    alignment=TA_LEFT, spaceBefore=0, spaceAfter=0,
                )
                cell_style_h = ParagraphStyle(
                    "tablecellh", parent=cell_style, fontSize=10, leading=15,
                    alignment=TA_CENTER,
                )
                data = [[Paragraph("故事元素", cell_style_h),
                         Paragraph("概念对应", cell_style_h)]]
                for row in fable["table_rows"]:
                    cells = (row + ["", ""])[:2]
                    data.append([
                        Paragraph(_escape_xml(cells[0]), cell_style),
                        Paragraph(_escape_xml(cells[1]), cell_style),
                    ])
                tbl = Table(data, colWidths=[7.5 * cm, 8.0 * cm])
                tbl.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#D9D9D9")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#999999")),
                    ("LEFTPADDING", (0, 0), (-1, -1), 4),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                    ("TOPPADDING", (0, 0), (-1, -1), 3),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                ]))
                story.append(tbl)
            else:
                story.append(Paragraph("（无）", s_quote))

            if fable["source_line"]:
                story.append(Paragraph(f"📝 {_escape_xml(fable['source_line'])}", s_source))
            story.append(Paragraph("· · ·", s_sep))
            story.append(Spacer(1, 12))

        # ----- 原文 part -----
        chapter_sections = sections_by_chapter.get(slug, [])
        if chapter_sections:
            story.append(Spacer(1, 80))
            s_divider = ParagraphStyle("divider", fontName=cjk_name, fontSize=13,
                                       leading=22, alignment=TA_CENTER,
                                       textColor=colors.HexColor("#666666"))
            story.append(Paragraph("—— 原文 ——", s_divider))
            story.append(PageBreak())

            s_source_h3 = ParagraphStyle("sourceh3", fontName=cjk_name, fontSize=13,
                                         leading=20, spaceBefore=12, spaceAfter=6, bold=True)
            s_source_body = ParagraphStyle("sourcebody", fontName=cjk_name, fontSize=10,
                                           leading=17, firstLineIndent=20, spaceAfter=3)
            s_source_sep = ParagraphStyle("sourcesep", fontName=cjk_name, fontSize=9,
                                          leading=12, alignment=TA_CENTER,
                                          textColor=colors.HexColor("#999999"),
                                          spaceBefore=4, spaceAfter=4)

            for sec in chapter_sections:
                for para in re.split(r"\n\s*\n", sec["body"]):
                    para = para.strip()
                    if not para:
                        continue
                    if re.fullmatch(r"-{3,}", para):
                        story.append(Paragraph("· · ·", s_source_sep))
                        continue
                    if para.startswith("## "):
                        story.append(Paragraph(_escape_xml(para[3:].strip()), s_source_h3))
                        continue
                    if para.startswith("!["):
                        # Image — skip in print
                        continue
                    story.append(Paragraph(_escape_xml(para).replace("\n", "<br/>"),
                                           s_source_body))

    doc.build(story)


# ---------------------------------------------------------------------------
# Orchestrator
# ---------------------------------------------------------------------------

def main():
    print(f"[book] fable root: {FABLES}")
    fables_by_chapter = {}
    sections_by_chapter = {}
    for slug, title in CHAPTERS:
        files = fable_files_in_chapter(slug)
        if not files:
            raise RuntimeError(f"No fables found for chapter {slug}")
        fables_by_chapter[slug] = []
        for p in files:
            fable = parse_fable(p)
            m = re.match(r'^(\d+)-', p.stem)
            sec_no = int(m.group(1)) if m else 0
            fables_by_chapter[slug].append((sec_no, fable, p))
        print(f"  chapter '{slug}': {len(files)} fables")

        # Load sections for this chapter
        sec_files = section_files_in_chapter(slug)
        sections_by_chapter[slug] = [parse_section(p) for p in sec_files]
        print(f"  chapter '{slug}': {len(sec_files)} source sections")

    # Backwards-compatible: also stash sections under '<slug>_sections' so the
    # markdown writer can find them via the same dict.
    for slug, _ in CHAPTERS:
        fables_by_chapter[f"{slug}_sections"] = sections_by_chapter[slug]

    print("[md] writing merged markdown...")
    md_lines = write_merged_markdown(fables_by_chapter)
    print(f"  -> {OUT_MD}  ({md_lines} lines, {OUT_MD.stat().st_size} bytes)")

    # Remove the temporary <slug>_sections keys before passing to docx/pdf
    # (they each have their own way to access sections; passing them would
    # confuse the existing iteration).
    fables_only = {slug: fables_by_chapter[slug] for slug, _ in CHAPTERS}

    print("[docx] building Word document...")
    build_docx(fables_only, sections_by_chapter, cjk_name="PingFang SC")
    print(f"  -> {OUT_DOCX}  ({OUT_DOCX.stat().st_size} bytes)")

    print("[pdf] building PDF...")
    cjk = _register_cjk_font()
    build_pdf(fables_only, sections_by_chapter, cjk_name=cjk)
    print(f"  -> {OUT_PDF}  ({OUT_PDF.stat().st_size} bytes)")

    print("[book] done.")


if __name__ == "__main__":
    main()
